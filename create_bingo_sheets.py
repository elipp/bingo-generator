from argparse import ArgumentParser
from pathlib import Path
from random import shuffle

from docx import Document


parser = ArgumentParser(
    description="Generate random bingo sheets from a list of bingo texts."
)
parser.add_argument(
    "-t",
    "--template",
    default="wedding_bingo_template.docx",
    help="File containing a bingo sheet template.",
)
parser.add_argument(
    "-n",
    "--number-of-sheets",
    dest="num_sheets",
    type=int,
    default="50",
    help="Number of bingo sheets to generate.",
)
parser.add_argument(
    "-b",
    "--bingo-texts",
    default="bingo_texts.txt",
    help="Text file containing bingo table texts.",
)
parser.add_argument(
    "-o", "--output-folder", default="./output", help="Output folder for bingo sheets."
)

args = parser.parse_args()

doc = Document(args.template)

with open(args.bingo_texts, "r", encoding="utf-8") as f:
    texts = [l.strip() for l in f.readlines()]

Path(args.output_folder).mkdir(parents=True, exist_ok=True)

table = doc.tables[0]
for i in range(args.num_sheets):
    shuffle(texts)
    text_rows = [texts[j : j + 5] for j in range(0, 25, 5)]

    for table_row, text_row in zip(table.rows, text_rows):
        for cell, text in zip(table_row.cells, text_row):
            cell.paragraphs[0].runs[0].text = text

    doc.save(f"{args.output_folder}/page{i}.docx")
